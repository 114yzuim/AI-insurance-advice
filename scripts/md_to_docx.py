"""
將 paper_draft_v1.md 轉換為格式化的 .docx（Pandoc/LaTeX 不可用環境下的替代方案）。
處理：# 標題層級、**粗體**、Markdown 表格、``` code block（等寬字體）、---分隔線（分頁）。
不處理：複雜巢狀清單、行內連結轉超連結（保留為純文字 URL）。

執行：python scripts/md_to_docx.py
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = Path(__file__).parent.parent / "paper_draft_v1.md"
OUT = Path(__file__).parent.parent / "paper_final.docx"

BODY_FONT = "Times New Roman"
CJK_FONT = "Source Han Serif TC" if False else "Microsoft JhengHei"  # fallback CJK font
MONO_FONT = "Consolas"


def set_cjk_font(run, font_name=CJK_FONT):
    run.font.name = BODY_FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def add_runs_with_bold(paragraph, text, base_size=11, mono=False):
    """Parse **bold** and *italic* markers within a line and add runs."""
    # Split on bold (**...**) first, then italic (*...*) within non-bold parts.
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith('**') and part.endswith('**')
        if is_bold:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(base_size)
            if mono:
                run.font.name = MONO_FONT
            else:
                set_cjk_font(run)
            continue
        # within non-bold text, split on single-asterisk italics
        subparts = re.split(r'(?<!\*)\*([^*\n]+?)\*(?!\*)', part)
        # re.split with one capture group alternates [text, italic, text, italic, ...]
        for idx, sub in enumerate(subparts):
            if not sub:
                continue
            is_italic = idx % 2 == 1
            run = paragraph.add_run(sub)
            run.italic = is_italic
            run.font.size = Pt(base_size)
            if mono:
                run.font.name = MONO_FONT
            else:
                set_cjk_font(run)


def parse_table(lines, start_idx):
    """Parse a markdown table starting at start_idx. Returns (rows, next_idx)."""
    rows = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith('|'):
        line = lines[i].strip()
        if re.match(r'^\|[\s:|-]+\|$', line):
            i += 1
            continue  # separator row
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
        i += 1
    return rows, i


def strip_md_inline(text):
    """Remove markdown bold markers for cell text length calc, keep for run parsing."""
    return text


def main():
    text = SRC.read_text(encoding='utf-8')
    lines = text.split('\n')

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = BODY_FONT
    style.font.size = Pt(11)

    # Page margins ~1 inch
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    i = 0
    in_code_block = False
    code_buffer = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code block toggle
        if stripped.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_buffer = []
            else:
                in_code_block = False
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                for cline in code_buffer:
                    run = p.add_run(cline + '\n')
                    run.font.name = MONO_FONT
                    run.font.size = Pt(9.5)
                # shading-ish: just keep plain, leave background to user
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Horizontal rule -> page break (skip for first occurrence)
        if stripped == '---':
            doc.add_page_break()
            i += 1
            continue

        # Headers
        m = re.match(r'^(#{1,3})\s+(.*)$', stripped)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2)
            h = doc.add_heading(level=level if level <= 3 else 3)
            add_runs_with_bold(h, heading_text, base_size=18 - level * 2)
            i += 1
            continue

        # Table
        if stripped.startswith('|'):
            rows, next_i = parse_table(lines, i)
            if rows:
                n_cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=n_cols)
                table.style = 'Light Grid Accent 1'
                for r_idx, row in enumerate(rows):
                    for c_idx in range(n_cols):
                        cell_text = row[c_idx] if c_idx < len(row) else ''
                        cell = table.cell(r_idx, c_idx)
                        cell.paragraphs[0].clear() if cell.paragraphs[0].runs else None
                        add_runs_with_bold(cell.paragraphs[0], cell_text, base_size=9.5)
                doc.add_paragraph()
            i = next_i
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Italic-only line markers like *text*
        # Regular paragraph
        p = doc.add_paragraph()
        add_runs_with_bold(p, stripped, base_size=11)
        i += 1

    doc.save(OUT)
    print(f"完成！已存至 {OUT}")


if __name__ == "__main__":
    main()
